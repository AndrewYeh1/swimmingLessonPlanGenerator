import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches

from templates import lessonPlan


class Word:
    def __init__(self, plan: lessonPlan):
        self.plan = plan

    def export(self):
        for day in self.plan.dayList:
            self.exportDay(day)

    def exportDay(self, day):
        doc = docx.Document()
        # creates table
        table = doc.add_table(rows=3, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # resizes table
        widths = (Inches(1), Inches(4), Inches(1), Inches(1))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # generate header
        # course
        courseCell = table.cell(0, 0).merge(table.cell(0, 1))
        courseCell.text = self.plan.course
        # wsi
        wsiCell = table.cell(0, 2).merge(table.cell(0, 3))
        wsiCell.text = "WSI:", self.plan.wsi
        # lesson num
        lessonNumCell = table.cell(1, 0)
        lessonNumCell.text = "Lesson #" + str(1)
        # date time
        dateTimeCell = table.cell(1, 1)
        dateTimeCell.text = "Date: \nTime: "
        # location
        locationCell = table.cell(1, 2).merge(table.cell(1, 3))
        locationCell.text = "Location:", self.plan.location
        # headers for activities
        table.cell(2, 0).paragraphs[0].add_run("Time").bold = True
        table.cell(2, 1).paragraphs[0].add_run("Activity").bold = True
        table.cell(2, 2).paragraphs[0].add_run("Formation").bold = True
        table.cell(2, 3).paragraphs[0].add_run("Equipment").bold = True

        for i in day:
            table = self.addActivity(table, i)

        # exports the word document
        doc.save("test.docx")

    # noinspection PyMethodMayBeStatic
    def addActivity(self, table, activity):
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(str(activity.time) + " min")
        row.cells[1].paragraphs[0].add_run("\n".join(activity.description))
        row.cells[2].paragraphs[0].add_run("Formation")
        row.cells[3].paragraphs[0].add_run("Equipment")
        return table

    def setFilePath(self):
        pass

    def setList(self):
        pass
